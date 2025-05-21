from docx import Document
from docx.shared import Pt
import pandas as pd
from datetime import datetime
import os

# Carrega os dados da planilha e do modelo de notificação
df = pd.read_excel("planilhas/criar_notificacoes.xlsx")
grupos = df.groupby("cod_cliente")
modelo_path = "modelo/Modelo_notificacao.docx"

# Variáveis fixas do ambiente teste 
SUA_CIDADE = os.getenv("SUA_CIDADE", "cidade-sp")
AGENCIA = os.getenv("AGENCIA", "0001-X")
ENDERECO_AGENCIA = os.getenv("ENDERECO_AGENCIA", "Rua da Agência, 123")
CNPJ = os.getenv("CNPJ", "00.000.000/0000-00")

# Função auxiliar para extrair o primeiro nome do cliente
def extrair_primeiro_nome(nome):
    return nome.split()[0].capitalize()

# Gera documentos por cliente
for cod_cliente, grupo in grupos:
    dados = grupo.iloc[0]
    doc = Document(modelo_path)

    # Substituições simples
    substituicoes = {
        "LOCAL": SUA_CIDADE,
        "DATA": datetime.today().strftime("%d/%m/%Y"),
        "CLIENTE": dados.get("nome_cliente", ""),
        "ENDERECO_CLIENTE": dados.get("endereco_cliente", ""),
        "CIDADE_CLIENTE": dados.get("cidade_cliente", ""),
        "UF_CLIENTE": dados.get("uf_cliente", ""),
        "CEP_CLIENTE": dados.get("cep_cliente", ""),
        "AGENCIA": AGENCIA,
        "ENDERECO_AGENCIA": ENDERECO_AGENCIA,
        "CNPJ": CNPJ
    }

    # Substitui os campos fixos no corpo, lembrando que os campos que estão no dict substituicoes, também estão no modelo de notificação 
    for par in doc.paragraphs:
        for chave, valor in substituicoes.items():
            if f"«{chave}»" in par.text:
                for run in par.runs:
                    run.text = run.text.replace(f"«{chave}»", str(valor))

    # Localiza e substitui o marcador de operações, para incluir lista depois 
    for i, par in enumerate(doc.paragraphs):
        if "«OPERACOES_BLOCO»" in par.text:
            par.clear()

            # Cabeçalho com alinhamento monoespaçado
            cabecalho = f"{'Produto'.ljust(30)}{'Operação'.ljust(25)}Vencimento"
            run = par.add_run(cabecalho)
            run.font.name = 'Courier New'
            run.font.size = Pt(11)
            run.bold = True

            # Adiciona as operações abaixo do cabeçalho
            for op in grupo.itertuples():
                produto = str(op.produto)[:20]
                operacao = str(op.operacao)[:20]
                vencimento = pd.to_datetime(op.vencimento).strftime("%d/%m/%Y")
                linha_texto = f"{produto.ljust(30)}{operacao.ljust(25)}{vencimento}"

                novo_par = doc.add_paragraph()
                novo_run = novo_par.add_run(linha_texto)
                novo_run.font.name = 'Courier New'
                novo_run.font.size = Pt(11)

                par._element.addnext(novo_par._element)
            break  # Só processa o primeiro marcador

    # Salva o documento final
    primeiro_nome = extrair_primeiro_nome(dados["nome_cliente"])
    data_hoje = datetime.today().strftime("%Y-%m-%d")
    nome_arquivo = f"notificacoes/notificacao_{primeiro_nome}_{cod_cliente}_{data_hoje}.docx"
    doc.save(nome_arquivo)
